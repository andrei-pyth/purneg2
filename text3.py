import data as cr
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor

class DocText():

    def __init__(self, comp):
        self.document = Document()
        self.comp = comp

        section = self.document.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        
        
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Time New Roman'
        self.font.size = Pt(14)
        self.font.bold = False

        self.style2 = self.document.styles['Heading 1']
        self.style2.next_paragraph_style = self.document.styles['Normal']
        self.font2 = self.style2.font
        self.font2.color.rgb = RGBColor(80, 80, 80)
        self.font2.name = 'Arial'
        self.font2.size = Pt(18)
        self.font2.italic = True
        self.font2.bold = False
        
        self.front_page()
        self.common_info()

    def front_page(self):
        self.paragraph = self.document.add_paragraph('Результаты проверки компании \n\n\n\n\n')
        self.paragraph.style = self.document.styles['Heading 1']
        self.paragraph.alignment = 1

        self.sentence = self.paragraph.add_run(self.comp.full_name_text[0] + '\n' + self.comp.full_name_text[1] + '\n\n\n\n\n')
        self.sentence.font.bold = True
        self.sentence.font.size = Pt(20)

        self.sentence = self.paragraph.add_run('на предмет соответствия требованиям\n к контрагентам\n\n')
        self.sentence.font.bold = False

        self.sentence = self.paragraph.add_run('"' + self.comp.customer.upper() + '"\n\n\n\n\n\n')
        self.sentence.font.bold = True

        self.sentence = self.paragraph.add_run(self.comp.date)
        self.sentence.font.bold = False

        self.sentence = self.paragraph.add_run('\nг.Санкт-Петербург')

        self.document.add_page_break()

    def common_info(self):

        self.document.add_paragraph('Общая информация:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.common_info_text + '\n').alignment = 3
        
        self.document.add_paragraph('Полное наименование компании:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.full_name + '\n')

        self.document.add_paragraph('Краткое наименование компании:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.short_name + '\n')
                
        self.document.add_paragraph('Налоговый режим:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.srn + '\n')
        
        self.document.add_paragraph('Форма собственности:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.okfs_text + '\n')
        
        self.document.add_paragraph('Организационно-правовая форма:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.okopf_text + '\n')

        self.document.add_paragraph('При регистрации организации присвоены:\n').style = self.document.styles['Heading 1']
        table = self.document.add_table(rows=9, cols=2)
        table.cell(0, 0).text = 'ОГРН'
        table.cell(0, 1).text = self.comp.ogrn
        table.cell(1, 0).text = 'ИНН'
        table.cell(1, 1).text = self.comp.inn
        table.cell(2, 0).text = 'КПП'
        if self.comp.kpp:
            table.cell(2, 1).text = self.comp.kpp
        else:
            table.cell(2, 1).text = ''
        table.cell(3, 0).text = 'ОКПО'
        table.cell(3, 1).text = self.comp.okpo
        table.cell(4, 0).text = 'ОКФС'
        table.cell(4, 1).text = self.comp.okfs
        table.cell(5, 0).text = 'ОКОПФ'
        table.cell(5, 1).text = self.comp.okopf
        table.cell(6, 0).text = 'ОКАТО'
        table.cell(6, 1).text = self.comp.okato
        table.cell(7, 0).text = 'ОКТМО'
        table.cell(7, 1).text = self.comp.oktmo
        table.cell(8, 0).text = 'ОКОГУ'
        table.cell(8, 1).text = self.comp.okogy
        
        self.document.add_paragraph('Статус:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.status + '\n')

        self.document.add_paragraph('Дата первичной регистрации:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.reg_date + '\n')

        self.document.add_paragraph('Размер уставного капитала:').style = self.document.styles['Heading 1']
        self.document.add_paragraph(f'\n\t{self.comp.capital.amount} рублей\n')

        self.document.add_page_break()

        self.document.add_paragraph('Юридический адрес компании:').style = self.document.styles['Heading 1']
        self.document.add_paragraph(self.comp.location_adress.address.map_photo) 
        try:
            lst1 = [self.comp.location_adress.address.index, self.comp.location_adress.address.city, self.comp.location_adress.address.street, self.comp.location_adress.address.building, self.comp.location_adress.address.appartment[0]]
        except AttributeError:
            lst1 = self.comp.location_adress.address
        for item1 in lst1:
            try:
                self.document.add_paragraph(f'\n\t{item1}')
            except IndexError:
                print('Error in formating occured =====> {item1}')
                self.document.add_paragraph(self.comp.location_adress.address.address)
                continue
        
        self.document.add_paragraph('Официальный сайт и интернет-ресурсы компании:\n\n').style = self.document.styles['Heading 1']

        self.document.add_paragraph('Филиалы и представительства:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.filials.filials + '\n')
        
        self.document.add_paragraph('Учредители компании:').style = self.document.styles['Heading 1']
        for count, item in enumerate(self.comp.founders.persons):
            self.document.add_paragraph(f'\n\t{count+1}) {item.surname.capitalize()}, {item.name.capitalize()}, {item.paternal.capitalize()} \n\tИНН: {item.inn}\n\tc {item.reg_date}')
        
        self.document.add_paragraph('Руководители компании:').style = self.document.styles['Heading 1']
        self.document.add_paragraph(f'\n\t{self.comp.chefs.post.post_name} (c {self.comp.chefs.post.reg_date})\n\t{self.comp.chefs.person.surname.capitalize()} {self.comp.chefs.person.name.capitalize()} {self.comp.chefs.person.paternal.capitalize()}\n\tИНН: {self.comp.chefs.person.inn}\n')

        self.document.add_paragraph('Последние изменения в ЕГРЮЛ:').style = self.document.styles['Heading 1']
        try:
            self.document.add_paragraph(f'\n{self.comp.egrul_changes[-1].date} - {self.comp.egrul_changes[-1].reason}\n\n{self.comp.egrul_changes[-2].date} - {self.comp.egrul_changes[-2].reason}\n\n{self.comp.egrul_changes[-3].date} - {self.comp.egrul_changes[-3].reason}\n')
        except IndexError:
            for item2 in self.comp.egrul_changes:
                print(f'{item2.date} ====> {item2.reason}\n\n\n')
            print('IndexError')

        
        self.document.add_paragraph('Опубликованные юридически значимые факты:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.important_facts + '\n')
        
        self.document.add_paragraph('Данные о текущей деятельности:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.current_text + '\n')

        self.document.add_paragraph('Сведения об уплаченных организацией налогах и сборах:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.taxes + '\n')
        
        self.document.add_paragraph('Сведения о заблокированных счетах:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.accounts_blocked + '\n')
        
        self.document.add_paragraph('Данные о внешнеэкономической деятельности:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.ved + '\n')
        
        self.document.add_paragraph('Виды деятельности: ').style = self.document.styles['Heading 1']
        for count, item in enumerate(self.comp.okveds):
            self.document.add_paragraph(f'\n{self.comp.okveds[count].number}    {self.comp.okveds[count].name};')

        self.document.add_paragraph('Участие в государственных закупках или госконтрактах:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.gos_zakupka + '\n' + self.comp.gos_kontrakt + '\n')
        
        self.document.add_paragraph('Данные об имеющихся лицензиях:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.lisences + '\n')

        self.document.add_paragraph('Данные о проверке благонадёжности:').style = self.document.styles['Heading 1']
        self.document.add_paragraph('\n\t' + self.comp.reliability + '\n')
        
        self.document.add_page_break()
        
        self.document.add_paragraph('Иные значимые факторы: ').style = self.document.styles['Heading 1']
        self.document.add_paragraph(f'\n{self.comp.facts}\n')
        
        self.document.add_page_break()
        
        self.document.add_paragraph('Приложения: ').style = self.document.styles['Heading 1']
        self.document.add_paragraph(f'\n{self.comp.applications}\n')
        
        self.document.add_page_break()
        


def save_doc(doc):
    fl = input('Введите название файла с результатом: ')
    doc.save(fl)

def main():
    comp = cr.Company()
    doc = DocText(comp)
    save_doc(doc.document)

if __name__ == '__main__':
    main()
      
